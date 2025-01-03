import os
from docx import Document


def replace_text_in_docx(file_path, output_path, old_text, new_text,old_text1,new_text1):
    try:
        # 打开文档
        doc = Document(file_path)

        # 遍历文档中的每个段落
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if old_text in run.text:
                    # 替换文本
                    run.text = run.text.replace(old_text, new_text)
                if old_text1 in run.text:
                    run.text = run.text.replace(old_text1, new_text1)

        # 遍历文档中的每个表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if old_text in run.text:
                                # 替换文本
                                run.text = run.text.replace(old_text, new_text)

        # 保存修改后的文档到目标文件夹
        doc.save(output_path)
        print(f"文件已处理并保存: {output_path}")
    except FileNotFoundError:
        print(f"文件未找到: {file_path}")
    except Exception as e:
        print(f"处理文件时出错: {file_path}, 错误: {e}")


def process_folder(input_folder, output_folder, old_text, new_text, old_text1, new_text1):
    # 确保目标文件夹存在
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # 遍历输入文件夹中的所有 .docx 文件
    for root, dirs, files in os.walk(input_folder):
        for file in files:
            if file.endswith('.doc'):
                input_file_path = os.path.join(root, file)
                output_file_path = os.path.join(output_folder, file)
                print(f"处理文件: {input_file_path}")
                replace_text_in_docx(input_file_path, output_file_path, old_text, new_text,old_text1, new_text1)


def calculate():
    # 示例调用
    input_folder = r'E:\JXdata\Python\wan\ArbitrationGet\文书下载'
    output_folder = r'E:\JXdata\Python\wan\ArbitrationGet\文书下载_处理后'
    process_folder(input_folder, output_folder, '绵阳', '凉山', '绵仲裁', '凉仲裁')


if __name__ == '__main__':
    calculate()