
# #
#
# #
# # import re
# # from docx import Document
# # import os
# #
# # # Word文档路径
# # doc_path = r'D:\work\WanxiangWork\xl\ABC安全B20220414之2021版中华人民共和国安全生产法释义全文.docx'
# #
# # # 读取Word文档
# # doc = Document(doc_path)
# #
# # # 初始化内容字符串
# # content = ""
# #
# # # 逐行读取文档内容
# # for para in doc.paragraphs:
# #     content += para.text + "\n"
# #
# # pattern = r"(?s)(?<=【条文主旨】).*?(?=【条文主旨】)"
# #
# # # 使用正则表达式分割内容
# # matches = re.finditer(pattern, content)
# #
# # # 初始化分割后的列表
# # sections = []
# #
# # # 收集分割后的段落
# # for match in matches:
# #     # 添加匹配到的内容
# #     sections.append(match.group())
# #
# #
# # text_o=[]
# # # 打印每个分割出来的部分
# # for i, section in enumerate(sections):
# #     html_text="【条文主旨】"+"\n"+section.strip()
# #     html_clear=html_text.split('\n')
# #     text_o.append(html_clear)
# #
# # for num,j in enumerate(text_o):
# #     old_html=j[-1]
# #     del j[-1]
# #
# #     # 检查新的 j 列表的最后一项是否包含 "第.*?条"
# #     if j:  # 再次确认 j 不为空
# #         last_element = j[-1]
# #         pattern = r"第.*?条"
# #
# #         if re.search(pattern, last_element):
# #             del j[-1]
# #         if j:
# #             last_element_text = j[-1]
# #             pattern = r"第.*?条"
# #
# #             if re.search(pattern, last_element):
# #                 del j[-1]
# #
# #         # 将处理后的 j 写入文件
# #         output_file = os.path.join(r'D:\work\下载分割', f"section_{num + 1}.txt")
# #         with open(output_file, 'w', encoding='utf-8') as f:
# #             for line in j:
# #                 f.write(line + "\n")
# #                 print('下载好了',f'section_{num + 1}.txt')
# #
# #     print("All sections have been processed and saved.")
#
#
#
#
#
#
#
# import os
# import re
#
# # 指定输入文件夹路径
# input_folder = r'D:\work\下载分割'
# # 指定输出文件夹路径
# output_folder = r'D:\work\下载分割处理后的新文件'
#
# # 创建输出文件夹
# os.makedirs(output_folder, exist_ok=True)
#
# # 遍历文件夹中的所有 .txt 文件
# for filename in os.listdir(input_folder):
#     if filename.endswith('.txt'):
#         input_path = os.path.join(input_folder, filename)
#         output_path = os.path.join(output_folder, filename)
#
#         # 逐行读取并处理文件
#         text_all = "<p>\n"
#         with open(input_path, 'r', encoding='utf-8') as file:
#             for line in file:
#                 line = "　　"+line+"　　"+"<br/>\n"
#                 text_all += line
#         text_all += "</p>\n"
#
#         # 将处理后的结果写入新文件
#         with open(output_path, 'w', encoding='utf-8') as file:
#             file.write(text_all)
#
#         print(f"Processed and saved: {output_path}")
#
# print("All files have been processed and saved.")
from asyncore import write










import pyodbc
import re
from docx import Document
import os

def save_sql_BidDocument(sql):
    '''
    用于插入数据库
    :param sql
    :return:
    '''
    connect = pyodbc.connect(
        'DRIVER={SQL Server};SERVER=47.97.3.24,14333;DATABASE=ourdata;UID=saa;PWD=1+2-3..*Qwe!@#;'
        'charset=gbk')
    # 创建游标对象
    cursor = connect.cursor()
    # sql = "INSERT INTO [自收录数据].dbo.[专项补充收录] ([唯一标志],[法规标题],[全文],[发布部门],[类别],[发布日期],[效力级别],[实施日期],[发文字号],[时效性],[来源],[收录时间]) VALUES (?,?,?,?,?,?,?,?,?,?,?,?)"
    # cursor.executemany(sql, all_result)
    cursor.execute(sql)

    cursor.commit()
    cursor.close()
    connect.close()



# def read_all(old_html):
#
#     num=int(838883424)
#     for i in range(13,15):
#         wenj_path = f"D:\work\下载分割处理后的新文件\section_{i}.txt"
#         with open(wenj_path, 'r', encoding='utf-8') as file:
#             read_cpntent  = file.read()
#             num+=1
#             write_to_oa(read_cpntent,num,old_html)
#

def read_all():
    # Word文档路径
    doc_path = r'D:\work\WanxiangWork\xl\ABC安全B20220414之2021版中华人民共和国安全生产法释义全文.docx'

    # 读取Word文档
    doc = Document(doc_path)

    # 初始化内容字符串
    content = ""

    # 逐行读取文档内容
    for para in doc.paragraphs:
        content += para.text + "\n"

    pattern = r"(?s)(?<=【条文主旨】).*?(?=【条文主旨】)"

    # 使用正则表达式分割内容
    matches = re.finditer(pattern, content)

    # 初始化分割后的列表
    sections = []

    # 收集分割后的段落
    for match in matches:
        # 添加匹配到的内容
        sections.append(match.group())

    text_o = []
    old_html_list=[]
    numty = int(838883424)
    # 打印每个分割出来的部分
    for i, section in enumerate(sections):
        html_text = "【条文主旨】" + "\n" + section.strip()
        html_clear = html_text.split('\n')
        text_o.append(html_clear)

    for num, j in enumerate(text_o):
        if num >= 11:
            old_html = j[-1]
            old_html_list.append(old_html)


    for j in old_html_list:
        # wenj_path = f"D:\work\下载分割处理后的新文件\section_{i}.txt"
        # with open(wenj_path, 'r', encoding='utf-8') as file:
        #     read_cpntent  = file.read()
            numty+=1
            write_to_oa(numty,j)





def write_to_oa(num,old_html):
    cgid = '820886279'
    name='《中华人民共和国安全生产法》（2021修订）释义'
    tgid='805306584'
    chapternum='1'
    ctitle='第一章　总　　则'
    gid='330429'
    sql = rf"INSERT INTO [ourdata].dbo.[syitemour] ([唯一标志],[cgid],[name],[tgid],[chapternum],[ctitle],[gid], [itemcontent]) VALUES ('{num}','{cgid}','{name}','{tgid}','{chapternum}','{ctitle}','{gid}','{old_html}')"
    save_sql_BidDocument(sql)
    print(sql)
    print("---插入一条了")


if __name__ == '__main__':
    read_all()
    # itemcontent()


